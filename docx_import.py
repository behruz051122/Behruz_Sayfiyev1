# docx_import.py
# Word (.docx) fayldan savollarni ommaviy o'qib olish.
#
# KUTILADIGAN SHABLON (har bir savol shu formatda yoziladi):
#
#   1. Savol matni shu yerga yoziladi?
#   (agar kerak bo'lsa, aynan shu yerga — savol bilan variantlar orasiga —
#   rasm/jadval/grafik joylashtiring: Word'da "Ilova -> Rasm" yoki
#   "Ilova -> Jadval" orqali)
#   A) birinchi variant
#   B) ikkinchi variant
#   C) uchinchi variant
#   D) to'rtinchi variant
#   Javob: B
#
# Savol raqami "1." yoki "1)" bilan, variantlar "A)" yoki "A." bilan
# boshlanishi kifoya (katta-kichik harf farqi muhim emas). "To'g'ri javob-B"
# yoki "Javob:B" kabi variantlar ham qo'llab-quvvatlanadi. Agar 4 ta variant
# bitta qatorda (masalan jadval ustunlaridek tab bilan ajratilgan) yozilgan
# bo'lsa, bu ham avtomatik ajratib olinadi. Word jadvallari (haqiqiy
# "Insert -> Table" orqali qo'yilgan) ham o'qiladi va ASL katak-katak
# ko'rinishida (matn shaklida emas) ilovada ko'rsatiladi.
#
# QULAY USUL — YASHIL RANG BILAN JAVOB BELGILASH:
# "Javob: B" qatorini har safar qo'lda yozish o'rniga, to'g'ri variantning
# MATNINI (yoki harfini) Word'da YASHIL rangga bo'yab qo'ysangiz ham bo'ladi
# — dastur buni avtomatik tanib, to'g'ri javob sifatida oladi. Agar bitta
# savolda ham yashil belgilash, ham "Javob:" qatori bo'lsa — YASHIL rang
# ustuvor hisoblanadi (chunki "Javob:" qatori ko'pincha shablondan
# nusxalanib, xato holda bir xil harfda qolib ketishi mumkin).

import io
import json
import re

QUESTION_RE = re.compile(r"^\s*(\d+)[\.\)]\s*(.*)$")
OPTION_RE = re.compile(r"^\s*[A-Da-d][\.\)]")
MULTI_OPTION_RE = re.compile(r"([A-Da-d])[\.\)]\s*")
# ANSWER_RE endi qatorning BOSHIDA emas, ISTALGAN joyida "javob" so'zini
# qidiradi ("To'g'ri javob-B", "Javob: B", "javob - b" barchasi mos keladi) —
# faqat qator harf bilan tugashi kerak.
ANSWER_RE = re.compile(r"javob\s*[:\-]?\s*([A-Da-d])\s*$", re.IGNORECASE)

BLIP_TAG = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
REL_EMBED_ATTR = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"

# Word'da rasm qo'yilganda ko'pincha ishlatiladigan format(lar) — <img> tegida
# to'g'ridan-to'g'ri ko'rsatib bo'ladiganlari. EMF/WMF kabi vektor formatlar
# brauzerda ochilmaydi, shu sababli ularni o'tkazib yuboramiz (savol matni va
# variantlar baribir to'g'ri qo'shiladi, faqat rasmsiz).
IMAGE_EXT_BY_CONTENT_TYPE = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/jpg": ".jpg",
    "image/gif": ".gif",
    "image/webp": ".webp",
}


def _extract_images(paragraph, document_part):
    """Paragraf ichidagi rasmlarni (bytes, content_type) ro'yxati sifatida qaytaradi."""
    images = []
    for blip in paragraph._p.findall(".//" + BLIP_TAG):
        rId = blip.get(REL_EMBED_ATTR)
        if not rId:
            continue
        try:
            part = document_part.related_parts[rId]
            images.append((part.blob, part.content_type))
        except KeyError:
            continue
    return images


def _is_greenish_rgb(rgb) -> bool:
    """Rang YASHIL ohangdami (00B050, 92D050 va shunga o'xshashlar)."""
    if rgb is None:
        return False
    try:
        r, g, b = rgb[0], rgb[1], rgb[2]
    except Exception:
        return False
    return g > 100 and g > r + 30 and g > b + 30


def _is_marker_rgb(rgb) -> bool:
    """Rang TO'G'RI JAVOBNI BELGILASH uchun ishlatilganmi.

    NEGA FAQAT YASHIL EMAS: amalda o'qituvchilar bitta hujjat ichida ham
    turli rang ishlatishadi — bir savolda yashil, boshqasida sariq. Faqat
    yashilni tan olish 50 savollik faylning 25 tasini "javobi yo'q" deb
    o'tkazib yuborishga olib kelgan edi.

    Shuning uchun endi: qora/oq/kulrang bo'lmagan, ya'ni ATAYLAB
    tanlangan HAR QANDAY ajratuvchi rang belgilash deb qabul qilinadi.
    Bu yashil, sariq, pushti, ko'k — barchasini qamrab oladi."""
    if rgb is None:
        return False
    try:
        r, g, b = rgb[0], rgb[1], rgb[2]
    except Exception:
        return False
    # Oq yoki juda och (deyarli oq) — bo'yoq emas
    if r > 235 and g > 235 and b > 235:
        return False
    # Qora yoki juda to'q — bu odatda oddiy matn rangi
    if r < 60 and g < 60 and b < 60:
        return False
    # Kulrang (kanallar deyarli teng) — bo'yoq emas
    if max(r, g, b) - min(r, g, b) < 25:
        return False
    return True


def _load_green_styles(document):
    """Ba'zi hujjatlarda to'g'ri javob to'g'ridan-to'g'ri matn/fon rangini
    o'zgartirish orqali emas, balki maxsus ravishda YASHIL FONLI qilib
    QAYTA TA'RIFLANGAN PARAGRAF/BELGI USLUBI (masalan "Heading 1",
    "Heading 2") orqali belgilanadi — o'qituvchi to'g'ri javob qatoriga
    shu uslubni qo'llaydi (masalan Word asboblar panelidan "Sarlavha 1"ni
    tanlaydi), uslubning o'zi esa fayl ichida (word/styles.xml) yashil fon
    bilan qayta ta'riflab qo'yilgan bo'ladi. Bunday holda variant matnini
    o'z ichiga olgan RUN'da HECH QANDAY to'g'ridan-to'g'ri formatlash
    bo'lmaydi (rang butunlay USLUBDAN meros olinadi) — shuning uchun buni
    alohida, styles.xml'ni oldindan o'qib chiqib aniqlash kerak.
    Natija: yashil fonli/rangli uslublarning styleId to'plami
    (masalan {"Heading1", "Heading2", "Heading1Char", "Heading2Char"})."""
    from docx.oxml.ns import qn

    green_style_ids = set()
    try:
        styles_element = document.styles.element
    except Exception:
        return green_style_ids
    if styles_element is None:
        return green_style_ids

    for style in styles_element.findall(qn("w:style")):
        style_id = style.get(qn("w:styleId"))
        if not style_id:
            continue
        rPr = style.find(qn("w:rPr"))
        if rPr is None:
            continue
        is_green = False

        shd = rPr.find(qn("w:shd"))
        if shd is not None:
            fill = shd.get(qn("w:fill"))
            if fill and fill.upper() not in ("AUTO", "FFFFFF"):
                try:
                    rgb = tuple(int(fill[i:i + 2], 16) for i in (0, 2, 4))
                    if _is_greenish_rgb(rgb):
                        is_green = True
                except Exception:
                    pass

        if not is_green:
            color_el = rPr.find(qn("w:color"))
            if color_el is not None:
                val = color_el.get(qn("w:val"))
                if val and val.upper() != "AUTO":
                    try:
                        rgb = tuple(int(val[i:i + 2], 16) for i in (0, 2, 4))
                        if _is_greenish_rgb(rgb):
                            is_green = True
                    except Exception:
                        pass

        if not is_green:
            highlight_el = rPr.find(qn("w:highlight"))
            if highlight_el is not None:
                val = (highlight_el.get(qn("w:val")) or "").upper()
                if val and val not in ("NONE", "WHITE", "AUTO"):
                    is_green = True

        if is_green:
            green_style_ids.add(style_id)

    return green_style_ids


def _paragraph_style_id(paragraph):
    """Paragrafning bevosita bog'langan uslubi (pStyle) ID'sini qaytaradi
    (masalan "Heading1"), yoki bo'lmasa None."""
    from docx.oxml.ns import qn

    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return None
    return pStyle.get(qn("w:val"))


def _run_is_green(run, green_style_ids=None, paragraph_style_id=None) -> bool:
    """Run (matn bo'lagi) YASHIL rangda yozilganmi, YASHIL STANDART
    highlight bilan belgilanganmi, YASHIL FON (shading/rang bilan bo'yash)
    bilan belgilanganmi, YOKI yashil fonli qilib QAYTA TA'RIFLANGAN
    USLUBGA (masalan "Heading 1") ega ekanmi tekshiradi — Word'da "Matn
    ranggi", "Highlight", "Shading/Fon rangi" va "Paragraf/Belgi uslubi"
    bir-biridan FARQLI XML mexanizmlari, o'qituvchilar esa turlicha usulda
    (ba'zan hatto uslub tanlash orqali) belgilaydi — shuning uchun
    barchasi tekshiriladi."""
    try:
        color = run.font.color
        if color is not None and color.type is not None and color.rgb is not None:
            if _is_greenish_rgb(color.rgb):
                return True
    except Exception:
        pass
    try:
        hl = run.font.highlight_color
        if hl is not None and str(hl).upper() not in ("NONE", "WHITE", "AUTO"):
            # Har qanday marker rangi (yashil, sariq, pushti...) — belgilash
            return True
    except Exception:
        pass
    try:
        from docx.oxml.ns import qn
        rPr = run._r.find(qn("w:rPr"))
        if rPr is not None:
            shd = rPr.find(qn("w:shd"))
            if shd is not None:
                fill = shd.get(qn("w:fill"))
                if fill and fill.upper() not in ("AUTO", "FFFFFF"):
                    rgb = tuple(int(fill[i:i + 2], 16) for i in (0, 2, 4))
                    if _is_marker_rgb(rgb):
                        return True
    except Exception:
        pass

    if green_style_ids:
        # Run o'zining alohida BELGI uslubiga (rStyle) ega bo'lishi mumkin —
        # bu paragraf uslubidan ustuvor bo'ladi (Word'dagi meros tartibi).
        try:
            from docx.oxml.ns import qn
            rPr = run._r.find(qn("w:rPr"))
            if rPr is not None:
                rStyle = rPr.find(qn("w:rStyle"))
                if rStyle is not None:
                    r_style_id = rStyle.get(qn("w:val"))
                    if r_style_id and r_style_id in green_style_ids:
                        return True
        except Exception:
            pass
        # Aks holda, run RUN o'ziga xos rStyle'ga ega bo'lmasa, paragrafning
        # UMUMIY uslubidan (pStyle) meros oladi.
        if paragraph_style_id and paragraph_style_id in green_style_ids:
            return True

    return False


def _load_numbering_starts(document):
    """Ba'zi Word hujjatlarida variant harflari ("C)", "D)" kabi) qo'lda
    yozilmagan, balki Word'ning AVTOMATIK ro'yxat raqamlash funksiyasi
    orqali chizilgan bo'ladi — bunday holda paragraph.text ICHIDA harf
    umuman ko'rinmaydi (Word uni faqat EKRANDA chizadi, fayl matnida
    saqlamaydi). Buni to'g'ri tiklash uchun word/numbering.xml faylidan har
    bir ro'yxat (numId) qanday harfdan boshlanishini (masalan "C" dan,
    ya'ni 3-harfdan) va formatini (katta/kichik harf) o'qib olamiz.
    Natija: {numId: {"start": int, "numFmt": str}}"""
    from docx.oxml.ns import qn

    result = {}
    try:
        numbering_part = document.part.numbering_part
    except Exception:
        return result
    if numbering_part is None:
        return result
    root = numbering_part.element

    abstract_info = {}
    for abstract_num in root.findall(qn("w:abstractNum")):
        abs_id = abstract_num.get(qn("w:abstractNumId"))
        lvl0 = None
        for lvl in abstract_num.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) == "0":
                lvl0 = lvl
                break
        if lvl0 is None:
            continue
        start_el = lvl0.find(qn("w:start"))
        fmt_el = lvl0.find(qn("w:numFmt"))
        abstract_info[abs_id] = {
            "start": int(start_el.get(qn("w:val"))) if start_el is not None else 1,
            "numFmt": fmt_el.get(qn("w:val")) if fmt_el is not None else "decimal",
        }

    for num in root.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        abs_ref = num.find(qn("w:abstractNumId"))
        if abs_ref is None:
            continue
        abs_id = abs_ref.get(qn("w:val"))
        if abs_id in abstract_info:
            result[num_id] = abstract_info[abs_id]
    return result


def _format_numbering_value(value: int, num_fmt: str) -> str:
    """Ro'yxat qiymatini (masalan 3) mos formatga o'giradi (masalan
    upperLetter uchun "C)")."""
    if num_fmt == "upperLetter":
        return f"{chr(ord('A') + (value - 1) % 26)})"
    if num_fmt == "lowerLetter":
        return f"{chr(ord('a') + (value - 1) % 26)})"
    return f"{value})"


def _resolve_auto_list_prefix(paragraph, numbering_info, numid_counters):
    """Berilgan paragraf Word'ning avtomatik ro'yxat raqamlashidan
    foydalansa (ilvl=0), tegishli "C)" kabi belgini hisoblab qaytaradi.
    Aks holda bo'sh satr qaytaradi. numid_counters — har bir numId qaysi
    qiymatgacha yetganini hujjat davomida kuzatib boradi (chunki bitta
    ro'yxat bir necha paragrafdan iborat bo'lishi mumkin: masalan bitta
    numId ostida ketma-ket C, D kelishi kerak)."""
    from docx.oxml.ns import qn

    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return ""
    num_pr = pPr.find(qn("w:numPr"))
    if num_pr is None:
        return ""
    ilvl_el = num_pr.find(qn("w:ilvl"))
    if ilvl_el is not None and ilvl_el.get(qn("w:val")) != "0":
        return ""
    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        return ""
    num_id = num_id_el.get(qn("w:val"))
    info = numbering_info.get(num_id)
    if not info:
        return ""
    current_value = numid_counters.get(num_id, info["start"] - 1) + 1
    numid_counters[num_id] = current_value
    return _format_numbering_value(current_value, info["numFmt"]) + " "


def _split_options_with_color(paragraph, prefix="", green_style_ids=None):
    """Bitta qatorda bitta YOKI bir nechta variant bo'lishi mumkin (masalan
    "A) 0   B) 25   C) 50   D) 75" — jadval ustunlaridek tab bilan ajratilgan).
    Paragraf RUN'lari (Word'dagi matn formatlash bo'laklari) orqali har bir
    variantning matnini VA shu variant YASHIL rang bilan belgilanganmi-
    yo'qmi aniqlaydi. Natija: {harf: {"text": ..., "green": bool}}"""
    paragraph_style_id = _paragraph_style_id(paragraph) if green_style_ids else None

    runs_info = []
    pos = len(prefix)
    full_text = prefix
    for run in paragraph.runs:
        t = run.text
        runs_info.append((pos, pos + len(t), run))
        full_text += t
        pos += len(t)

    matches = list(MULTI_OPTION_RE.finditer(full_text))
    if not matches:
        return {}

    result = {}
    for i, m in enumerate(matches):
        letter = m.group(1).upper()
        content_start = m.end()
        content_end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
        content = full_text[content_start:content_end].strip(" \t")

        # Ranglash tekshiruvi harf belgisi ("B)")dan boshlab olinadi — chunki
        # ba'zi o'qituvchilar faqat harfni, ba'zilari butun variant matnini
        # yashil rangga bo'yaydi, ikkalasi ham hisobga olinishi kerak.
        color_start, color_end = m.start(), content_end
        is_green = any(
            _run_is_green(run, green_style_ids, paragraph_style_id)
            for r_start, r_end, run in runs_info
            if r_start < color_end and r_end > color_start
        )
        # Agar shu qatorda umuman RUN topilmasa (masalan butun matn faqat
        # avtomatik ro'yxat prefiksidan iborat bo'lsa) ham, paragraf
        # o'zining USLUBI orqali (masalan "Heading 1") yashil bo'lishi
        # mumkin — bunday holatni ham hisobga olamiz.
        if not is_green and green_style_ids and paragraph_style_id in green_style_ids:
            is_green = True

        result[letter] = {"text": content, "green": is_green}
    return result


def _table_rows(table):
    """Haqiqiy Word jadvalini (Insert -> Table orqali qo'yilgan) qatorlar
    ro'yxati sifatida qaytaradi — matnga aylantirmasdan, ASL katak-katak
    tuzilishini saqlab qolish uchun (frontend buni haqiqiy <table> sifatida
    chizadi)."""
    rows = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        rows.append(cells)
    return rows


def _table_as_single_option(table, green_style_ids=None):
    """Ba'zan (odatda tasodifiy formatlash natijasida) bitta variantning
    matni alohida ma'lumot jadvali o'rniga Word jadvaliga tushib qolgan
    bo'ladi (masalan "D) ..." matni). Bunday holatni jadvalning BARCHA
    matnini birlashtirib, natija "A) matn" uslubiga mos kelsa aniqlaymiz —
    aks holda bu oddiy ma'lumot jadvali (davriy jadval, izotoplar foizi
    va h.k.) deb hisoblanadi. Mos kelsa (harf, matn, yashilmi) qaytaradi,
    aks holda None."""
    from docx.oxml.ns import qn

    seen_texts = []
    is_green = False
    for row in table.rows:
        for cell in row.cells:
            txt = cell.text.strip()
            if txt and txt not in seen_texts:
                seen_texts.append(txt)
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is not None:
                    shd = tcPr.find(qn("w:shd"))
                    if shd is not None:
                        fill = shd.get(qn("w:fill"))
                        if fill and fill.upper() not in ("AUTO", "FFFFFF"):
                            try:
                                rgb = tuple(int(fill[i:i + 2], 16) for i in (0, 2, 4))
                                if _is_greenish_rgb(rgb):
                                    is_green = True
                            except Exception:
                                pass
                if not is_green and green_style_ids:
                    # Katak ichidagi paragraf(lar) ham "Heading 1" kabi
                    # yashil fonli uslubga ega bo'lishi mumkin — shu ham
                    # tekshiriladi (jadval katakchalari uchun ham xuddi
                    # oddiy paragraflardagidek uslub merosxo'rligi ishlaydi).
                    for cell_paragraph in cell.paragraphs:
                        p_style_id = _paragraph_style_id(cell_paragraph)
                        if p_style_id and p_style_id in green_style_ids:
                            is_green = True
                            break
                        for run in cell_paragraph.runs:
                            if _run_is_green(run, green_style_ids, p_style_id):
                                is_green = True
                                break
                        if is_green:
                            break

    joined = " ".join(seen_texts).strip()
    m = MULTI_OPTION_RE.match(joined)
    if not m:
        return None
    letter = m.group(1).upper()
    content = joined[m.end():].strip()
    return letter, content, is_green


def _iter_block_items(document):
    """python-docx standart holda paragraflar (doc.paragraphs) va jadvallar
    (doc.tables) ni ALOHIDA, hujjatdagi HAQIQIY tartibini yo'qotib ro'yxat
    qiladi. Savol qaysi jadvalga tegishli ekanini bilish uchun bizga aniq
    hujjat tartibi kerak — shu funksiya har ikkalasini asl tartibda beradi."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


class ParsedQuestion:
    def __init__(self, order_num):
        self.order_num = order_num
        self.question_text = ""
        self.options = {}
        self.correct_letter = None   # "Javob: B" kabi MATN orqali ko'rsatilgan javob
        self.green_letter = None     # variant matni YASHIL rangda belgilangan bo'lsa
        self.image_bytes = None
        self.image_content_type = None
        self.table_data = None       # {"tables": [{"rows": [[...], ...]}, ...]}
        self.has_table = False

    @property
    def effective_correct_letter(self):
        """YASHIL rang bilan belgilash — "Javob:" qatoridan KO'RA ustuvor,
        chunki matn qatori ko'pincha shablondan nusxalanib, xato holda bir
        xil harfda qolib ketishi mumkin (o'qituvchining haqiqiy niyati odatda
        aynan YASHIL rangda ko'rinadi)."""
        return self.green_letter or self.correct_letter

    def is_complete(self) -> bool:
        letter = self.effective_correct_letter
        return (
            bool(self.question_text.strip())
            and len(self.options) == 4
            and letter in self.options
        )

    def missing_summary(self) -> str:
        problems = []
        if not self.question_text.strip():
            problems.append("savol matni yo'q")
        missing_opts = [l for l in "ABCD" if l not in self.options]
        if missing_opts:
            problems.append(f"variant(lar) yo'q: {', '.join(missing_opts)}")
        letter = self.effective_correct_letter
        if not letter:
            problems.append("to'g'ri javob ko'rsatilmagan (variantni YASHIL rangda belgilang yoki \"Javob: B\" kabi qator qo'shing)")
        elif letter not in self.options:
            problems.append(f"to'g'ri javob ({letter}) hech qanday variantga mos kelmadi")
        return "; ".join(problems) if problems else ""

    def table_data_json(self):
        return json.dumps(self.table_data, ensure_ascii=False) if self.table_data else None


def parse_docx(file_bytes: bytes):
    """
    .docx fayl baytlarini o'qib, ParsedQuestion obyektlari ro'yxatini qaytaradi.
    Har bir savol to'liq (is_complete()==True) yoki muammoli bo'lishi mumkin —
    chaqiruvchi (router) muammolilarni admin uchun aniq ko'rsatishi kerak.
    """
    from docx import Document
    from docx.table import Table

    doc = Document(io.BytesIO(file_bytes))
    document_part = doc.part

    # Ba'zi hujjatlarda variant harflari ("C)", "D)") Word'ning AVTOMATIK
    # ro'yxat raqamlashi orqali chiziladi va paragraph.text ichida umuman
    # ko'rinmaydi (masalan "A) 23,3" qo'lda yozilgan, keyingi "C) 69,9" esa
    # Word ro'yxati bo'lib, faylda faqat "69,9" saqlanadi). Buni oldindan
    # word/numbering.xml'dan o'qib, har bir shunday paragrafga to'g'ri
    # harfni qo'lda "tiklab" beramiz.
    numbering_info = _load_numbering_starts(doc)
    numid_counters = {}

    # Ba'zi hujjatlarda to'g'ri javob YASHIL rangga to'g'ridan-to'g'ri
    # bo'yalmagan, balki maxsus yashil fonli qilib qayta ta'riflangan
    # PARAGRAF USLUBI (masalan "Heading 1") orqali belgilanadi — buni
    # oldindan word/styles.xml'dan bilib olamiz (batafsili izoh
    # _load_green_styles() ichida).
    green_style_ids = _load_green_styles(doc)

    questions = []
    current = None
    order_counter = 0
    # Savol matni ICHIDA "1) ...", "2) ..." kabi raqamlangan bandlar (masalan
    # "quyidagi ifodalardan qaysilari to'g'ri?" turidagi savollarda) bo'lishi
    # mumkin — bularni HAQIQIY yangi savol bilan adashtirmaslik uchun, raqam
    # faqat OLDINGI tasdiqlangan savol raqamidan KATTA bo'lsagina, uni yangi
    # savol deb hisoblaymiz ("aniq +1" talab qilinmaydi — real hujjatlarda
    # savol raqamlari orasida uzilish bo'lishi mumkin, masalan 11 dan keyin
    # to'g'ridan-to'g'ri 14 kelishi; lekin savol ICHIDAGI band deyarli
    # doim 1 dan qayta boshlanadi, shuning uchun "kattami" tekshiruvi
    # ikkalasini ham to'g'ri ajratadi).
    last_confirmed_number = None

    for block in _iter_block_items(doc):
        if isinstance(block, Table):
            if current is not None:
                option_from_table = _table_as_single_option(block, green_style_ids)
                if option_from_table and len(current.options) < 4:
                    letter, content, is_green = option_from_table
                    current.options[letter] = content
                    if is_green:
                        current.green_letter = letter
                else:
                    current.has_table = True
                    rows = _table_rows(block)
                    if rows:
                        if current.table_data is None:
                            current.table_data = {"tables": []}
                        current.table_data["tables"].append({"rows": rows})
            continue

        paragraph = block
        auto_prefix = _resolve_auto_list_prefix(paragraph, numbering_info, numid_counters)
        raw_text = paragraph.text.strip()
        text = (auto_prefix + raw_text).strip() if auto_prefix else raw_text

        q_match = QUESTION_RE.match(text) if text else None
        opt_match = OPTION_RE.match(text) if text else None
        ans_match = ANSWER_RE.search(text) if text else None

        is_new_question = False
        if q_match:
            found_num = int(q_match.group(1))
            if current is None or last_confirmed_number is None or found_num > last_confirmed_number:
                is_new_question = True

        if is_new_question:
            if current is not None:
                questions.append(current)
            order_counter += 1
            current = ParsedQuestion(order_counter)
            remainder = q_match.group(2).strip()
            current.question_text = remainder
            last_confirmed_number = int(q_match.group(1))

            # Savol raqami bilan BIR QATORDA (masalan "...necha? A) 23,3")
            # birinchi variant ham boshlanib qolgan bo'lishi mumkin — bunday
            # holda uni ajratib olib, savol matnini shu belgigacha qisqartiramiz.
            inline_match = MULTI_OPTION_RE.search(remainder)
            if inline_match:
                current.question_text = remainder[:inline_match.start()].strip()
                inline_found = _split_options_with_color(paragraph, auto_prefix, green_style_ids)
                for letter, info in inline_found.items():
                    current.options[letter] = info["text"]
                    if info["green"]:
                        current.green_letter = letter
        elif current is not None and (opt_match or MULTI_OPTION_RE.search(text)):
            # Variant belgisi qatorning BOSHIDA bo'lmasligi ham mumkin —
            # masalan oldingi izoh matni bilan bir qatorda kelib qolgan
            # bo'lishi mumkin ("...II-band. A) I-to'g'ri..."). Shunday
            # holatda belgidan OLDINGI qism savol matniga qo'shiladi.
            first_match = MULTI_OPTION_RE.search(text)
            leading_text = text[:first_match.start()].strip() if first_match else ""
            if leading_text and not current.options:
                current.question_text = (current.question_text + " " + leading_text).strip()
            found = _split_options_with_color(paragraph, auto_prefix, green_style_ids)
            for letter, info in found.items():
                current.options[letter] = info["text"]
                if info["green"]:
                    current.green_letter = letter
        elif current is not None and ans_match:
            current.correct_letter = ans_match.group(1).upper()
        elif current is not None and text and not current.options:
            # Variant/javob qatoriga mos kelmagan, bo'sh bo'lmagan qator —
            # ko'p qatorli savol matnining davomi deb hisoblanadi (variantlar
            # boshlangandan keyingi izohli qatorlarga tegilmaydi).
            current.question_text = (current.question_text + " " + text).strip()

        if current is not None and current.image_bytes is None:
            imgs = _extract_images(paragraph, document_part)
            if imgs:
                current.image_bytes, current.image_content_type = imgs[0]

    if current is not None:
        questions.append(current)

    return questions
